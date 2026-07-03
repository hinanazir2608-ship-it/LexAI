
import os

# ============================================================
# ENV VARIABLES — directly set karo yahan
# ============================================================
from dotenv import load_dotenv
load_dotenv()
import json
from datetime import datetime
from pathlib import Path
import io

import flask
from flask_login import (
    LoginManager, login_user, logout_user,
    login_required, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.shared import RGBColor

from models import db, User, ChatSession, ChatMessage, ContractComparison
from rag_pipeline import (
    process_uploaded_pdf, load_vector_store,
    rag_legal_qa, rag_risk_analysis,
    rag_plain_english_summary, rag_missing_clauses,
    rag_risk_score,
    rag_compare_contracts,        # new
    process_multiple_pdfs,        # new
    )
load_dotenv()

app = flask.Flask(__name__)  # templates/ aur static/ automatically detect ho jayenge

app.config["SECRET_KEY"]                     = os.getenv("SECRET_KEY", "change-me")
app.config["SQLALCHEMY_DATABASE_URI"]        = os.getenv("SQLALCHEMY_DATABASE_URI", "sqlite:///legal_assistant.db")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"]                  = os.getenv("UPLOAD_FOLDER", "uploads")
app.config["MAX_CONTENT_LENGTH"]             = 16 * 1024 * 1024

ALLOWED_EXTENSIONS = {"pdf"}

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"
login_manager.login_message = "Please log in to access LexAI."

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/")
def index():
    if current_user.is_authenticated:
        return flask.redirect(flask.url_for("dashboard"))
    return flask.render_template("index.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return flask.redirect(flask.url_for("dashboard"))
    if flask.request.method == "POST":
        full_name = flask.request.form.get("full_name", "").strip()
        email     = flask.request.form.get("email", "").strip().lower()
        phone     = flask.request.form.get("phone", "").strip()
        password  = flask.request.form.get("password", "")
        confirm   = flask.request.form.get("confirm_password", "")

        if not full_name or not email or not password:
            flask.flask("All required fields must be filled.", "error")
            return flask.render_template("register.html")
        if password != confirm:
            flask.flask("Passwords do not match.", "error")
            return flask.render_template("register.html")
        if len(password) < 8:
            flask.flask("Password must be at least 8 characters.", "error")
            return flask.render_template("register.html")
        if User.query.filter_by(email=email).first():
            flask.flask("An account with this email already exists.", "error")
            return flask.render_template("register.html")

        new_user = User(
            full_name=full_name, email=email, phone=phone,
            password_hash=generate_password_hash(password),
        )
        db.session.add(new_user)
        db.session.commit()
        login_user(new_user)
        flask.flask(f"Welcome, {full_name}!", "success")
        return flask.redirect(flask.url_for("dashboard"))
    return flask.render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return flask.redirect(flask.url_for("dashboard"))
    if flask.request.method == "POST":
        email    = flask.request.form.get("email", "").strip().lower()
        password = flask.request.form.get("password", "")
        remember = flask.request.form.get("remember", False)
        user = User.query.filter_by(email=email).first()
        if not user or not check_password_hash(user.password_hash, password):
            flask.flask("Invalid email or password.", "error")
            return flask.render_template("login.html")
        login_user(user, remember=bool(remember))
        return flask.redirect(flask.request.args.get("next") or flask.url_for("dashboard"))
    return flask.render_template("login.html")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    flask.flask("Logged out securely.", "info")
    return flask.redirect(flask.url_for("index"))

@app.route("/dashboard")
@login_required
def dashboard():
    sessions = ChatSession.query.filter_by(
        user_id=current_user.id
    ).order_by(ChatSession.updated_at.desc()).all()
    return flask.render_template("dashboard.html", user=current_user, sessions=sessions)

@app.route("/upload", methods=["POST"])
@login_required
def upload():
    if "pdf_file" not in flask.request.files:
        return flask.jsonify({"success": False, "error": "No file provided"}), 400
    file         = flask.request.files["pdf_file"]
    session_name = flask.request.form.get("session_name", "").strip()
    if file.filename == "":
        return flask.jsonify({"success": False, "error": "No file selected"}), 400
    if not allowed_file(file.filename):
        return flask.jsonify({"success": False, "error": "Only PDF files accepted"}), 400

    filename    = secure_filename(file.filename)
    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
    os.makedirs(user_folder, exist_ok=True)
    pdf_path    = os.path.join(user_folder, filename)
    file.save(pdf_path)

    try:
        result = process_uploaded_pdf(pdf_path, str(current_user.id), user_folder)
    except Exception as e:
        return flask.jsonify({"success": False, "error": str(e)}), 500

    if not session_name:
        session_name = Path(filename).stem.replace("_", " ").title()

    new_session = ChatSession(
        user_id=current_user.id, session_name=session_name,
        document_filename=filename, document_path=pdf_path,
        vector_store_path=result["vector_store_path"],
        text_path=result["text_path"],
    )
    db.session.add(new_session)
    db.session.commit()

    return flask.jsonify({
        "success": True, "session_id": new_session.id,
        "session_name": session_name, "filename": filename,
        "chunks": result["total_chunks"],
    })

@app.route("/chat/<int:session_id>")
@login_required
def chat_view(session_id):
    chat_session = ChatSession.query.filter_by(
        id=session_id, user_id=current_user.id).first_or_404()
    messages = ChatMessage.query.filter_by(
        session_id=session_id).order_by(ChatMessage.created_at.asc()).all()
    return flask.render_template("chat.html", user=current_user,
                           chat_session=chat_session, messages=messages)

@app.route("/chat/<int:session_id>/query", methods=["POST"])
@login_required
def chat_query(session_id):
    chat_session = ChatSession.query.filter_by(
        id=session_id, user_id=current_user.id).first_or_404()
    data       = flask.request.get_json()
    feature    = data.get("feature", "qa")
    user_query = data.get("query", "").strip()

    try:
        vector_store = load_vector_store(chat_session.vector_store_path)
    except Exception as e:
        return flask.jsonify({"success": False, "error": str(e)}), 500

    full_text = ""
    if chat_session.text_path and os.path.exists(chat_session.text_path):
        with open(chat_session.text_path, "r", encoding="utf-8") as f:
            full_text = f.read()

    if user_query:
        db.session.add(ChatMessage(session_id=session_id, role="user",
                                   content=user_query, feature_type=feature))

    try:
        if feature == "qa":
            if not user_query:
                return flask.jsonify({"success": False, "error": "Please enter a question"}), 400
            response_content = rag_legal_qa(user_query, vector_store)
        elif feature == "risk":
            response_content = json.dumps(rag_risk_analysis(vector_store, full_text))
        elif feature == "summary":
            response_content = rag_plain_english_summary(full_text)
        elif feature == "missing":
            response_content = json.dumps(rag_missing_clauses(full_text))
        else:
            return flask.jsonify({"success": False, "error": "Unknown feature"}), 400
    except Exception as e:
        return flask.jsonify({"success": False, "error": str(e)}), 500

    db.session.add(ChatMessage(session_id=session_id, role="assistant",
                               content=response_content, feature_type=feature))
    chat_session.updated_at = datetime.utcnow()
    db.session.commit()

    return flask.jsonify({"success": True, "response": response_content, "feature": feature})

@app.route("/download/<int:session_id>")
@login_required
def download_report(session_id):
    chat_session = ChatSession.query.filter_by(
        id=session_id, user_id=current_user.id).first_or_404()
    messages = ChatMessage.query.filter_by(
        session_id=session_id).order_by(ChatMessage.created_at.asc()).all()

    doc = DocxDocument()
    title = doc.add_heading("Legal Analysis Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x23, 0x4e)
    doc.add_paragraph(f"Document: {chat_session.document_filename}")
    doc.add_paragraph(f"Attorney: {current_user.full_name}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}")
    disc = doc.add_paragraph("DISCLAIMER: AI analysis only. Not legal advice.")
    disc.runs[0].font.italic = True

    for msg in messages:
        if msg.role == "user":
            p = doc.add_paragraph()
            run = p.add_run(f"Query ({msg.feature_type.upper()}):")
            run.bold = True
            doc.add_paragraph(msg.content)
        else:
            p = doc.add_paragraph()
            run = p.add_run("AI Response:")
            run.bold = True
            try:
                doc.add_paragraph(json.dumps(json.loads(msg.content), indent=2))
            except Exception:
                doc.add_paragraph(msg.content)
        doc.add_paragraph("─" * 40)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return flask.send_file(buffer, as_attachment=True,
                     download_name=f"{secure_filename(chat_session.session_name)}_report.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    if flask.request.method == "POST":
        current_user.full_name = flask.request.form.get("full_name", current_user.full_name).strip()
        current_user.phone     = flask.request.form.get("phone", current_user.phone).strip()
        new_pw = flask.request.form.get("new_password", "").strip()
        if new_pw:
            if len(new_pw) < 8:
                flask.flask("Password must be at least 8 characters.", "error")
                return flask.redirect(flask.url_for("profile"))
            current_user.password_hash = generate_password_hash(new_pw)
        db.session.commit()
        flask.flask("Profile updated!", "success")
        return flask.redirect(flask.url_for("profile"))
    return flask.render_template("profile.html", user=current_user)

# ============================================================
# CONTRACT COMPARISON ROUTES
# ============================================================
@app.route("/compare")
@login_required
def compare():
    comparisons = ContractComparison.query.filter_by(
        user_id=current_user.id
    ).order_by(ContractComparison.created_at.desc()).all()
    return flask.render_template("compare.html", user=current_user, comparisons=comparisons)
@app.route("/compare/run", methods=["POST"])
@login_required

def compare_run():
    """Process two PDFs and run comparison."""

    if "doc1" not in flask.request.files or "doc2" not in flask.request.files:
        return flask.jsonify({"success": False, "error": "Both documents required"}), 400

    doc1 = flask.request.files["doc1"]
    doc2 = flask.request.files["doc2"]
    session_name = flask.request.form.get("session_name", "Contract Comparison").strip()

    if doc1.filename == "" or doc2.filename == "":
        return flask.jsonify({"success": False, "error": "Both files must be selected"}), 400

    if not allowed_file(doc1.filename) or not allowed_file(doc2.filename):
        return flask.jsonify({"success": False, "error": "Only PDF files accepted"}), 400

    # Save both files
    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
    os.makedirs(user_folder, exist_ok=True)

    doc1_filename = secure_filename(doc1.filename)
    doc2_filename = secure_filename(doc2.filename)

    doc1_path = os.path.join(user_folder, f"cmp1_{doc1_filename}")
    doc2_path = os.path.join(user_folder, f"cmp2_{doc2_filename}")

    doc1.save(doc1_path)
    doc2.save(doc2_path)

    # Extract texts
    try:
        from rag_pipeline import extract_text_from_pdf
        text1 = extract_text_from_pdf(doc1_path)
        text2 = extract_text_from_pdf(doc2_path)
    except Exception as e:
        return flask.jsonify({"success": False, "error": f"PDF extraction failed: {str(e)}"}), 500

    # Save texts
    text1_path = os.path.join(user_folder, f"cmp1_text_{current_user.id}.txt")
    text2_path = os.path.join(user_folder, f"cmp2_text_{current_user.id}.txt")

    with open(text1_path, "w", encoding="utf-8") as f:
        f.write(text1)
    with open(text2_path, "w", encoding="utf-8") as f:
        f.write(text2)

    # Run comparison
    try:
        result = rag_compare_contracts(text1, text2)
    except Exception as e:
        return flask.jsonify({"success": False, "error": f"Comparison failed: {str(e)}"}), 500

    # Save to DB
    comparison = ContractComparison(
        user_id=current_user.id,
        session_name=session_name,
        doc1_filename=doc1_filename,
        doc1_path=doc1_path,
        doc1_text_path=text1_path,
        doc2_filename=doc2_filename,
        doc2_path=doc2_path,
        doc2_text_path=text2_path,
        comparison_result=json.dumps(result),
    )
    db.session.add(comparison)
    db.session.commit()

    return flask.jsonify({
        "success": True,
        "comparison_id": comparison.id,
        "result": result,
    })


@app.route("/compare/result/<int:comparison_id>")
@login_required
def compare_result(comparison_id):
    """View saved comparison result."""
    comparison = ContractComparison.query.filter_by(
        id=comparison_id, user_id=current_user.id).first_or_404()

    result = {}
    if comparison.comparison_result:
        try:
            result = json.loads(comparison.comparison_result)
        except Exception:
            result = {}

    return flask.render_template(
        "compare_result.html",
        user=current_user,
        comparison=comparison,
        result=result,
    )


# ============================================================
# MULTI-DOCUMENT UPLOAD ROUTE
# ============================================================

@app.route("/upload-multi", methods=["POST"])
@login_required
def upload_multi():
    """Process multiple PDFs into one combined session."""
    print("Upload multi route hit!")
    print(f"Files: {flask.request.files}")
    if "pdf_files" not in flask.request.files:
        return flask.jsonify({"success": False, "error": "No files provided"}), 400

    files = flask.request.files.getlist("pdf_files")
    session_name = flask.request.form.get("session_name", "").strip()

    if not files or all(f.filename == "" for f in files):
        return flask.jsonify({"success": False, "error": "No files selected"}), 400

    if len(files) > 5:
        return flask.jsonify({"success": False, "error": "Maximum 5 files allowed"}), 400

    # Validate all files
    for f in files:
        if not allowed_file(f.filename):
            return flask.jsonify({"success": False, "error": f"{f.filename} is not a PDF"}), 400

    # Save all files
    user_folder = os.path.join(app.config["UPLOAD_FOLDER"], str(current_user.id))
    os.makedirs(user_folder, exist_ok=True)

    saved_paths = []
    filenames = []

    for f in files:
        filename = secure_filename(f.filename)
        path = os.path.join(user_folder, filename)
        f.save(path)
        saved_paths.append(path)
        filenames.append(filename)

    # Process through multi-doc pipeline
    try:
        result = process_multiple_pdfs(saved_paths, str(current_user.id), user_folder)
    except Exception as e:
        return flask.jsonify({"success": False, "error": str(e)}), 500

    # Create session name
    if not session_name:
        session_name = f"Multi-Doc: {', '.join(filenames[:2])}"
        if len(filenames) > 2:
            session_name += f" +{len(filenames)-2} more"

    # Save as ChatSession
    new_session = ChatSession(
        user_id=current_user.id,
        session_name=session_name,
        document_filename=", ".join(filenames),
        document_path=", ".join(saved_paths),
        vector_store_path=result["vector_store_path"],
        text_path=result["text_path"],
    )
    db.session.add(new_session)
    db.session.commit()

    return flask.jsonify({
        "success": True,
        "session_id": new_session.id,
        "session_name": session_name,
        "doc_count": result["doc_count"],
        "total_chunks": result["total_chunks"],
    })
@app.route("/risk-dashboard/<int:session_id>")
@login_required
def risk_dashboard(session_id):
    """Risk Score Dashboard — visual contract health report."""

    chat_session = ChatSession.query.filter_by(
        id=session_id, user_id=current_user.id).first_or_404()

    # Load full text
    full_text = ""
    if chat_session.text_path and os.path.exists(chat_session.text_path):
        with open(chat_session.text_path, "r", encoding="utf-8") as f:
            full_text = f.read()

    if not full_text:
        flask.flask("Document text not found.", "error")
        return flask.redirect(flask.url_for("dashboard"))

    # Generate risk score
    try:
        risk_data = rag_risk_score(full_text)
    except Exception as e:
        flask.flask(f"Risk analysis failed: {str(e)}", "error")
        return flask.redirect(flask.url_for("chat_view", session_id=session_id))

    return flask.render_template(
        "risk_dashboard.html",
        user=current_user,
        chat_session=chat_session,
        risk=risk_data,
    )
with app.app_context():
    db.create_all()
    print("Database tables created successfully!")

if __name__ == "__main__":
    app.run(debug=True, port=5000)